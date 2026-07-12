"""Build the approved-candidate neutral industry report DOCX asset.

This script is intentionally deterministic in content and does not invoke shell
commands. Run it with python-docx available, then privacy-scrub and render the
result before registering the asset digest.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = "17365D"
BLUE = "1F4E78"
TEAL = "2F75B5"
PALE_BLUE = "EEF4F8"
MID_GREY = "667085"
LIGHT_GREY = "E7EAF0"
TEXT = "202632"
WHITE = "FFFFFF"


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def _set_table_borders(table, color: str = "C9D2DC", size: str = "6") -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)


def _set_run_font(run, east_asia: str = "Microsoft YaHei", latin: str = "Arial") -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    _set_run_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    paragraph = normal.paragraph_format
    paragraph.space_after = Pt(6)
    paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.line_spacing = 1.12

    heading_specs = {
        "Heading 1": (18, NAVY, 18, 8),
        "Heading 2": (13.5, BLUE, 14, 6),
        "Heading 3": (11.5, TEAL, 10, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = document.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = document.styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string(MID_GREY)


def _configure_section(section) -> None:
    # Named override for China-facing business documents: A4 portrait.
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True


def _configure_header_footer(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("行业报告编写数字员工  /  临时中性模板")
    _set_run_font(run)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)
    paragraph.paragraph_format.space_after = Pt(3)
    bottom = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "8")
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), TEAL)
    bottom.append(border)
    paragraph._p.get_or_add_pPr().append(bottom)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Cm(15.92))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(11.5)
    table.columns[1].width = Cm(4.42)
    left = table.cell(0, 0).paragraphs[0]
    run = left.add_run("内部资料｜仅向任务委派人交付")
    _set_run_font(run)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)
    _add_page_number(table.cell(0, 1).paragraphs[0])


def _add_accent_bar(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(5)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), TEAL)
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def _add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(0.35)
    table.columns[1].width = Cm(15.57)
    _set_cell_shading(table.cell(0, 0), TEAL)
    content = table.cell(0, 1)
    _set_cell_shading(content, PALE_BLUE)
    _set_cell_margins(content, top=140, bottom=140, start=180, end=180)
    paragraph = content.paragraphs[0]
    run = paragraph.add_run(title)
    _set_run_font(run)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    run.font.size = Pt(10.5)
    paragraph.add_run("\n")
    body_run = paragraph.add_run(body)
    _set_run_font(body_run)
    body_run.font.size = Pt(9.5)
    paragraph.paragraph_format.space_after = Pt(0)


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        _set_run_font(run)
        paragraph.paragraph_format.space_after = Pt(3)


def _add_matrix(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    for index, (header, width) in enumerate(zip(headers, widths, strict=True)):
        cell = table.rows[0].cells[index]
        cell.width = Cm(width)
        _set_cell_shading(cell, NAVY)
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        _set_run_font(run)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(WHITE)
    _set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, (value, width) in enumerate(zip(values, widths, strict=True)):
            cells[column_index].width = Cm(width)
            _set_cell_margins(cells[column_index])
            if row_index % 2:
                _set_cell_shading(cells[column_index], "F7F9FB")
            paragraph = cells[column_index].paragraphs[0]
            run = paragraph.add_run(value)
            _set_run_font(run)
            run.font.size = Pt(8.5)
        _prevent_row_split(table.rows[-1])


def _add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(48)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("行业报告编写数字员工")
    _set_run_font(run)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)

    _add_accent_bar(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(34)
    title.add_run("2025年中国证券行业\n发展研究报告")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("临时中性模板 · asset candidate v1.0.0")

    document.add_paragraph().paragraph_format.space_before = Pt(78)
    info = document.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    info.columns[0].width = Cm(4.0)
    info.columns[1].width = Cm(8.0)
    values = [
        ("报告期", "2025年"),
        ("目标受众", "证券公司总经理"),
        ("编制部门", "战略发展部"),
        ("版本状态", "供任务委派人评审"),
    ]
    for row, (label, value) in zip(info.rows, values, strict=True):
        for cell in row.cells:
            _set_cell_margins(cell, top=130, bottom=130)
        _set_cell_shading(row.cells[0], PALE_BLUE)
        label_run = row.cells[0].paragraphs[0].add_run(label)
        _set_run_font(label_run)
        label_run.bold = True
        label_run.font.color.rgb = RGBColor.from_string(BLUE)
        value_run = row.cells[1].paragraphs[0].add_run(value)
        _set_run_font(value_run)
    _set_table_borders(info, color=LIGHT_GREY, size="4")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(68)
    run = paragraph.add_run("内部资料｜仅向任务委派人交付")
    _set_run_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def _add_template_guide(document: Document) -> None:
    document.add_heading("模板使用说明", level=1)
    document.add_paragraph(
        "本模板用于 v0.1.0 首个行业报告 Playbook 的 DOCX 生成与验收。"
        "它定义中性版式和章节骨架，不代表公司品牌规范；正式公司模板接入时应创建新的 asset_version。"
    )
    _add_callout(
        document,
        "正式产物规则",
        "文件名包含报告名、报告期、artifact version 与批准状态；报告正文、来源清单和证据清单必须绑定同一 WorkItem。",
    )
    document.add_heading("推荐目录", level=2)
    _add_bullets(
        document,
        [
            "执行摘要：结论先行，列出关键数字、主要变化、五矿证券可行动建议和重大不确定性。",
            "行业发展概况：规模、监管、市场结构、盈利驱动和 2025 年关键变化。",
            "证券公司经营差异：统一口径比较收入、利润、资本效率、业务结构和风险。",
            "五矿证券业务线对标：财富、投行、自营权益、自营固收、资管逐线分析。",
            "借鉴点与行动建议：明确优先级、适用条件、依赖、风险和验证指标。",
            "来源与证据附录：记录来源定位、检索时间、许可、快照状态和关键 Claim 绑定。",
        ],
    )
    document.add_heading("版式合同", level=2)
    _add_matrix(
        document,
        ["要素", "冻结候选值", "生成要求"],
        [
            ["页面", "A4 纵向；四边 2.54 cm", "正文可用宽度 15.92 cm"],
            ["字体", "中文 Microsoft YaHei；西文 Arial", "缺失时由 Word 做平台字体替代"],
            ["标题", "18 / 13.5 / 11.5 pt", "稳定 section_id 不写入可变章节序号"],
            ["正文", "10.5 pt；1.12 倍行距", "段后 6 pt；避免超密排版"],
            ["页眉页脚", "中性角色名、内部资料、页码", "封面隐藏常规页眉"],
            ["颜色", "深蓝、蓝、浅蓝、灰", "不使用公司 Logo 或未批准品牌色"],
        ],
        [2.2, 5.2, 8.5],
    )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_contents(document: Document) -> None:
    document.add_heading("目录", level=1)
    entries = [
        ("S01", "执行摘要"),
        ("S02", "行业发展概况"),
        ("S03", "证券公司经营差异"),
        ("S04", "五矿证券业务线对标"),
        ("S05", "借鉴点与行动建议"),
        ("S06", "风险与不确定性"),
        ("A01", "来源与证据附录"),
    ]
    for section_id, title in entries:
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Cm(2.3)
        table.columns[1].width = Cm(13.62)
        left, right = table.rows[0].cells
        _set_cell_margins(left, top=110, bottom=110)
        _set_cell_margins(right, top=110, bottom=110)
        code = left.paragraphs[0].add_run(section_id)
        _set_run_font(code)
        code.bold = True
        code.font.color.rgb = RGBColor.from_string(TEAL)
        name = right.paragraphs[0].add_run(title)
        _set_run_font(name)
        name.font.size = Pt(11)
        name.font.color.rgb = RGBColor.from_string(NAVY)
    document.add_paragraph(
        "正式生成器依据稳定 section_id 生成目录和页码；章节重排不改变评审锚点。"
    ).paragraph_format.space_before = Pt(16)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_report_skeleton(document: Document) -> None:
    document.add_heading("执行摘要", level=1)
    document.add_paragraph(
        "执行摘要应在两页以内回答四件事：行业发生了什么、同业差异在哪里、"
        "五矿证券各业务线可以借鉴什么、哪些结论仍需人工判断。关键数字必须可追溯到证据。"
    )
    _add_callout(
        document,
        "总经理阅读视角",
        "先给结论和行动含义，再给支撑事实；区分事实、推断、建议与待验证假设。",
    )
    document.add_heading("关键结论呈现", level=2)
    _add_matrix(
        document,
        ["结论主题", "应回答的问题", "证据要求", "对五矿证券的意义"],
        [
            ["行业概况", "规模、结构和盈利驱动如何变化", "至少两类独立来源", "判断资源配置方向"],
            ["同业差异", "领先与承压机构差异来自哪里", "统一口径公司数据", "识别可复制能力"],
            ["业务对标", "五条业务线分别处于什么位置", "分业务可比指标", "形成逐线行动建议"],
            ["风险", "哪些数字或判断存在冲突", "保留冲突证据", "设置决策边界"],
        ],
        [2.4, 4.2, 3.7, 5.6],
    )
    document.add_heading("行业发展概况", level=1)
    document.add_paragraph(
        "本章使用统一报告期和统计口径描述市场规模、监管政策、竞争格局、收入利润、资本约束及技术趋势。"
        "同比、排名和份额必须注明分母、币种、单位和时间点。"
    )
    document.add_heading("指标示例", level=2)
    _add_matrix(
        document,
        ["指标", "2024", "2025", "变化", "来源编号"],
        [
            ["行业营业收入", "按来源口径填列", "按来源口径填列", "同比与驱动拆解", "SRC-001"],
            ["行业净利润", "按来源口径填列", "按来源口径填列", "同比与集中度", "SRC-002"],
            ["净资本", "期末口径", "期末口径", "资本效率变化", "SRC-003"],
        ],
        [3.4, 2.6, 2.6, 4.0, 3.3],
    )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    document.add_heading("证券公司经营差异", level=1)
    document.add_paragraph(
        "比较必须先建立可比公司集、数据版本和口径映射。对绝对规模、增长、盈利质量、"
        "资本效率和业务结构分别排序，避免用单一指标定义领先。"
    )
    document.add_heading("同业比较表", level=2)
    _add_matrix(
        document,
        ["机构", "收入/利润", "ROE/资本效率", "业务结构", "主要差异", "证据"],
        [
            ["可比机构 A", "统一单位", "统一年度", "分业务占比", "事实与解释分列", "CLM-001"],
            ["可比机构 B", "统一单位", "统一年度", "分业务占比", "事实与解释分列", "CLM-002"],
            ["五矿证券", "统一单位", "统一年度", "分业务占比", "对标差距与优势", "CLM-003"],
        ],
        [2.4, 2.6, 2.8, 3.0, 3.5, 1.6],
    )
    document.add_heading("五矿证券业务线对标", level=1)
    _add_matrix(
        document,
        ["业务线", "对标维度", "可借鉴做法", "适用前提", "建议指标"],
        [
            ["财富", "客户、产品、渠道、投顾", "说明能力机制而非只列案例", "客群与渠道匹配", "AUM、客户活跃、产品保有"],
            ["投行", "行业覆盖、项目质量、协同", "识别专业化和区域化打法", "人员与项目储备", "收入、项目数、承销质量"],
            ["自营权益", "策略、回撤、资本占用", "区分方向与策略收益", "风险预算明确", "收益率、回撤、资本回报"],
            ["自营固收", "交易、做市、信用研究", "比较资金与风险效率", "授信与限额完备", "久期、利差、资本回报"],
            ["资管", "产品、投研、渠道、规模质量", "识别主动管理优势", "牌照与渠道协同", "净流入、业绩、留存"],
        ],
        [2.1, 3.1, 4.2, 3.1, 3.4],
    )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    document.add_heading("借鉴点与行动建议", level=1)
    document.add_paragraph(
        "每项建议必须说明业务价值、执行动作、责任边界、依赖条件、风险和验证指标。"
        "数字员工可以生成候选建议，但重大资源配置结论必须由委派人批准。"
    )
    _add_matrix(
        document,
        ["优先级", "建议", "依据", "依赖/风险", "验证方式"],
        [
            ["P0", "近期可验证动作", "关键 Claim 与 Evidence", "列明数据与组织条件", "30—90 天指标"],
            ["P1", "能力建设动作", "同业机制与内部差距", "列明投入与协同", "季度里程碑"],
            ["观察", "尚不足以决策的方向", "冲突或不完整证据", "明确补充研究", "下一次复核点"],
        ],
        [2.0, 3.6, 3.8, 3.7, 2.8],
    )
    document.add_heading("风险与不确定性", level=2)
    _add_bullets(
        document,
        [
            "口径风险：公司披露口径、合并范围或报告期不一致。",
            "来源风险：商业数据库许可限制摘录或无法形成稳定快照。",
            "推断风险：公开信息不足以支持因果判断，只能形成待验证假设。",
            "时效风险：报告生成后市场、监管或公司数据发生变化。",
        ],
    )
    document.add_heading("来源与证据附录", level=1)
    _add_matrix(
        document,
        ["来源编号", "来源与定位", "检索时间", "许可/快照", "支持的 Claim"],
        [
            ["SRC-001", "公开或授权来源的稳定 locator", "ISO 8601", "许可说明与 digest", "CLM-001"],
            ["SRC-002", "Gildata MCP 查询摘要", "ISO 8601", "数据版本与 digest", "CLM-002"],
            ["SRC-003", "委派人授权的内部材料", "ISO 8601", "受控快照与等级", "CLM-003"],
        ],
        [2.2, 4.8, 2.7, 3.6, 2.6],
    )
    _add_callout(
        document,
        "引用规则",
        "正文使用稳定来源编号；来源清单记录 locator、retrieved_at、许可、snapshot status 与 digest。禁止把 token、签名 URL 或宿主路径写入产物。",
    )


def build(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = "行业报告编写数字员工临时中性模板"
    document.core_properties.subject = "v0.1.0 M0 Word asset candidate"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = "Neutral template; no company branding."

    _configure_styles(document)
    for section in document.sections:
        _configure_section(section)
        _configure_header_footer(section)

    _add_cover(document)
    _add_contents(document)
    _add_template_guide(document)
    _add_report_skeleton(document)

    # Ask Word-compatible clients to refresh fields such as page numbers.
    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.output)


if __name__ == "__main__":
    main()
